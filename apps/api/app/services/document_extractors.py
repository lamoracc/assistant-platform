from dataclasses import dataclass
from io import BytesIO
import logging
from pathlib import Path
from pathlib import PurePosixPath
import re
from typing import Any
from urllib.parse import urljoin

try:
    from bs4 import BeautifulSoup, Tag
except ImportError:
    BeautifulSoup = None
    Tag = Any

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

from app.services.chunking import TextBlock
from app.services.file_router import route_file
from app.services.language import detect_language
from app.services.text_sanitizer import sanitize_text

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    metadata: dict
    blocks: list[TextBlock]


def extract_document(
    content: bytes,
    filename: str,
    content_type: str,
) -> ExtractedDocument:
    routed = route_file(Path(filename), content_type.split(";")[0])
    if routed.route == "html":
        return extract_html(content, filename)
    if routed.route == "pdf":
        return extract_pdf(content, filename)
    if routed.route == "word":
        return extract_word(content, filename)
    if routed.route == "text":
        if routed.extension == ".md":
            return extract_markdown(content, filename)
        return extract_text_document(content, filename, routed.extension)
    raise ValueError("Unsupported document type.")


def is_html_document(filename: str, content_type: str) -> bool:
    lowered = filename.lower()
    content_type = content_type.split(";")[0]
    return (
        content_type in {"text/html", "application/xhtml+xml"}
        or lowered.endswith(".html")
        or lowered.endswith(".htm")
    )


def extract_pdf(content: bytes, filename: str) -> ExtractedDocument:
    if PdfReader is None:
        raise RuntimeError("pypdf is required to extract PDF documents.")
    reader = PdfReader(BytesIO(content))
    blocks: list[TextBlock] = []
    page_texts: list[str] = []

    for page_index, page in enumerate(reader.pages):
        page_text = sanitize_text(page.extract_text() or "")
        page_texts.append(page_text)
        image_refs = _extract_pdf_page_images(page, page_index + 1)
        for paragraph in page_text.split("\n\n"):
            cleaned = sanitize_text(paragraph)
            if cleaned:
                blocks.append(
                    TextBlock(
                        text=cleaned,
                        metadata={
                            "page": page_index + 1,
                            "source": filename,
                            "image_refs": image_refs,
                        },
                    )
                )

    metadata = {
        "filename": filename,
        "page_count": len(reader.pages),
        "pdf_metadata": _clean_pdf_metadata(reader.metadata or {}),
        "page_image_refs": [
            {"page": index + 1, "images": _extract_pdf_page_images(page, index + 1)}
            for index, page in enumerate(reader.pages)
        ],
        "tables_extracted_as_text": True,
        "language": detect_language("\n\n".join(page_texts)),
        "document_type": "pdf",
    }
    return ExtractedDocument(text="\n\n".join(page_texts).strip(), metadata=metadata, blocks=blocks)


def extract_word(content: bytes, filename: str) -> ExtractedDocument:
    if filename.lower().endswith(".doc"):
        return extract_legacy_doc(content, filename)
    return extract_docx(content, filename)


def extract_docx(content: bytes, filename: str) -> ExtractedDocument:
    if DocxDocument is None:
        raise RuntimeError("python-docx is required to extract DOCX documents.")
    document = DocxDocument(BytesIO(content))
    blocks: list[TextBlock] = []
    text_parts: list[str] = []

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = sanitize_text(paragraph.text)
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        text_parts.append(text)
        blocks.append(
            TextBlock(
                text=text,
                heading=text if style_name.lower().startswith("heading") else None,
                metadata={
                    "paragraph": paragraph_index,
                    "style": style_name,
                    "source": filename,
                },
            )
        )

    core = document.core_properties
    metadata = {
        "filename": filename,
        "author": core.author,
        "category": core.category,
        "comments": core.comments,
        "created": core.created.isoformat() if core.created else None,
        "identifier": core.identifier,
        "keywords": core.keywords,
        "language": core.language,
        "last_modified_by": core.last_modified_by,
        "modified": core.modified.isoformat() if core.modified else None,
        "subject": core.subject,
        "title": core.title,
        "version": core.version,
        "paragraph_count": len(document.paragraphs),
        "document_type": "docx",
        "language": detect_language("\n\n".join(text_parts)),
    }
    return ExtractedDocument(text="\n\n".join(text_parts), metadata=metadata, blocks=blocks)


def extract_legacy_doc(content: bytes, filename: str) -> ExtractedDocument:
    text = _extract_printable_text(content)
    blocks = [
        TextBlock(text=paragraph, metadata={"source": filename, "document_type": "doc"})
        for paragraph in text.split("\n\n")
        if paragraph.strip()
    ]
    metadata = {
        "filename": filename,
        "document_type": "doc",
        "extraction_method": "legacy_doc_printable_text_fallback",
        "language": detect_language(text),
    }
    return ExtractedDocument(text=text, metadata=metadata, blocks=blocks)


def extract_html(content: bytes, filename: str) -> ExtractedDocument:
    if BeautifulSoup is None:
        raise RuntimeError("beautifulsoup4 is required to extract HTML documents.")
    soup = BeautifulSoup(content, "html.parser")
    title = sanitize_text(soup.title.get_text(" ", strip=True)) if soup.title else None
    breadcrumbs = _extract_breadcrumbs(soup)
    support_asset_refs = _extract_support_asset_refs(soup, filename)
    _remove_html_chrome(soup)

    layer = soup.select_one("div#Layer1")
    body = soup.body or soup
    main = layer if layer and _has_readable_text(layer) else body

    links = _extract_links(main, filename)
    image_refs = _extract_image_refs(main, filename)
    blocks, current_heading = _extract_html_blocks(main, filename, title)

    if not blocks and main is not body:
        links = _extract_links(body, filename)
        image_refs = _extract_image_refs(body, filename)
        blocks, current_heading = _extract_html_blocks(body, filename, title)

    text = "\n\n".join(block.text for block in blocks if block.text).strip()
    metadata = {
        "source_filename": filename,
        "title": title,
        "breadcrumbs": breadcrumbs,
        "links": links,
        "image_refs": image_refs,
        "support_asset_refs": support_asset_refs,
        "main_heading": current_heading,
        "document_type": "html",
        "language": detect_language(text),
    }
    return ExtractedDocument(text=text, metadata=metadata, blocks=blocks)


def extract_text_document(
    content: bytes,
    filename: str,
    extension: str,
) -> ExtractedDocument:
    text = sanitize_text(content.decode("utf-8", errors="ignore"))
    if not text:
        text = sanitize_text(content.decode("latin-1", errors="ignore"))

    blocks: list[TextBlock] = []
    current_heading: str | None = None
    for index, paragraph in enumerate(text.split("\n\n")):
        cleaned = sanitize_text(paragraph)
        if not cleaned:
            continue
        is_heading = extension == ".md" and cleaned.startswith("#")
        if is_heading:
            current_heading = cleaned.lstrip("#").strip()
            blocks.append(
                TextBlock(
                    text=current_heading,
                    heading=current_heading,
                    metadata={"source": filename, "style": "heading 1"},
                )
            )
            continue
        blocks.append(
            TextBlock(
                text=cleaned,
                heading=current_heading,
                metadata={"source": filename, "paragraph": index},
            )
        )

    metadata = {
        "filename": filename,
        "document_type": "markdown" if extension == ".md" else "text",
        "language": detect_language(text),
    }
    return ExtractedDocument(text=text, metadata=metadata, blocks=blocks)


def extract_markdown(content: bytes, filename: str) -> ExtractedDocument:
    raw_text = content.decode("utf-8", errors="ignore")
    if not raw_text.strip():
        raw_text = content.decode("latin-1", errors="ignore")

    front_matter, body = _split_markdown_front_matter(raw_text)
    metadata = _parse_front_matter(front_matter)
    title = sanitize_text(metadata.get("title")) or Path(filename).stem
    breadcrumbs = _normalize_breadcrumbs(metadata.get("breadcrumbs"))

    cleaned_lines = _clean_markdown_lines(body)
    blocks = _markdown_blocks_from_lines(cleaned_lines, filename, title)
    text = "\n\n".join(block.text for block in blocks if block.text).strip()

    metadata = {
        "filename": filename,
        "source_filename": filename,
        "topic_id": sanitize_text(metadata.get("topic_id")),
        "title": title,
        "relative_path": sanitize_text(metadata.get("relative_path")),
        "section_group": sanitize_text(metadata.get("section_group")),
        "breadcrumbs": breadcrumbs,
        "document_type": "markdown",
        "language": detect_language(text),
    }
    return ExtractedDocument(text=text, metadata=metadata, blocks=blocks)


def _clean_pdf_metadata(metadata: dict) -> dict:
    cleaned: dict[str, str] = {}
    for key, value in metadata.items():
        cleaned[str(key).lstrip("/")] = str(value)
    return cleaned


def _split_markdown_front_matter(raw_text: str) -> tuple[str, str]:
    normalized = raw_text.replace("\r\n", "\n")
    if not normalized.startswith("---\n"):
        return "", normalized

    lines = normalized.split("\n")
    for index in range(1, len(lines)):
        if lines[index].strip() == "---":
            return "\n".join(lines[1:index]), "\n".join(lines[index + 1 :])
    return "", normalized


def _parse_front_matter(front_matter: str) -> dict:
    metadata: dict[str, object] = {}
    current_key: str | None = None

    for raw_line in front_matter.splitlines():
        line = raw_line.rstrip()
        if not line.strip() or line.lstrip().startswith("#"):
            continue

        if line.startswith((" ", "\t")) and current_key:
            item = line.strip()
            if item.startswith("- "):
                metadata.setdefault(current_key, [])
                if isinstance(metadata[current_key], list):
                    metadata[current_key].append(_strip_yaml_quotes(item[2:]))
            continue

        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        current_key = sanitize_text(key)
        value = value.strip()
        if not value:
            metadata[current_key] = []
            continue
        if value.startswith("[") and value.endswith("]"):
            metadata[current_key] = [
                _strip_yaml_quotes(item.strip())
                for item in value[1:-1].split(",")
                if item.strip()
            ]
        else:
            metadata[current_key] = _strip_yaml_quotes(value)

    return metadata


def _strip_yaml_quotes(value: object) -> str:
    text = sanitize_text(str(value))
    if len(text) >= 2 and text[0] == text[-1] and text[0] in {"'", '"'}:
        return text[1:-1]
    return text


def _normalize_breadcrumbs(value: object) -> list[str]:
    if isinstance(value, list):
        return [sanitize_text(str(item)) for item in value if sanitize_text(str(item))]
    text = sanitize_text(str(value or ""))
    if not text:
        return []
    delimiter = ">" if ">" in text else "/"
    return [part.strip() for part in text.split(delimiter) if part.strip()]


def _clean_markdown_lines(body: str) -> list[str]:
    lines = body.replace("\r\n", "\n").split("\n")
    cleaned: list[str] = []
    index = 0

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = sanitize_text(line)

        if _is_markdown_image_only(stripped) or _is_breadcrumb_only_line(stripped):
            index += 1
            continue

        if _starts_markdown_table(lines, index):
            table_lines, next_index = _collect_markdown_table(lines, index)
            if _is_navigation_table(table_lines):
                index = next_index
                continue
            cleaned.extend(table_lines)
            index = next_index
            continue

        cleaned.append(line)
        index += 1

    return cleaned


def _markdown_blocks_from_lines(
    lines: list[str],
    filename: str,
    title: str,
) -> list[TextBlock]:
    blocks: list[TextBlock] = [
        TextBlock(
            text=title,
            heading=title,
            metadata={"source": filename, "style": "heading 1"},
        )
    ]
    paragraph: list[str] = []
    current_heading = title
    index = 0

    def flush_paragraph() -> None:
        nonlocal paragraph
        text = sanitize_text("\n".join(paragraph))
        paragraph = []
        if text:
            blocks.append(
                TextBlock(
                    text=text,
                    heading=current_heading,
                    metadata={"source": filename, "markdown_block": "paragraph"},
                )
            )

    while index < len(lines):
        line = lines[index]
        stripped = sanitize_text(line)

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        heading = _markdown_heading_text(stripped)
        if heading:
            flush_paragraph()
            current_heading = heading
            blocks.append(
                TextBlock(
                    text=heading,
                    heading=heading,
                    metadata={"source": filename, "style": "heading 1"},
                )
            )
            index += 1
            continue

        if _starts_markdown_table(lines, index):
            flush_paragraph()
            table_lines, next_index = _collect_markdown_table(lines, index)
            table_text = _markdown_table_to_text(table_lines)
            if table_text:
                blocks.append(
                    TextBlock(
                        text=table_text,
                        heading=current_heading,
                        metadata={
                            "source": filename,
                            "markdown_block": "table",
                            "html_tag": "table",
                        },
                    )
                )
            index = next_index
            continue

        paragraph.append(stripped)
        index += 1

    flush_paragraph()
    return blocks


def _is_markdown_image_only(line: str) -> bool:
    return bool(re.fullmatch(r"\s*!\[[^\]]*]\([^)]+\)\s*", line))


def _is_breadcrumb_only_line(line: str) -> bool:
    if ">" not in line:
        return False
    link_count = len(re.findall(r"\[[^\]]+]\([^)]+\)|\[[^\]]+]", line))
    plain = re.sub(r"\[[^\]]+]\([^)]+\)", "", line)
    plain = re.sub(r"\[[^\]]+]", "", plain)
    plain = plain.replace(">", "").strip()
    return link_count >= 2 and not plain


def _starts_markdown_table(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    return "|" in lines[index] and bool(re.search(r"^\s*\|?\s*:?-{3,}", lines[index + 1]))


def _collect_markdown_table(lines: list[str], index: int) -> tuple[list[str], int]:
    table: list[str] = []
    while index < len(lines) and "|" in lines[index]:
        table.append(lines[index].strip())
        index += 1
    return table, index


def _is_navigation_table(table_lines: list[str]) -> bool:
    text = " ".join(table_lines).lower()
    tokens = ["previous", "next", "printer", "print", "welcome", "contents"]
    meaningful = re.sub(r"[\|\-:\s\[\]\(\)_./]+", "", text)
    return any(token in text for token in tokens) and len(meaningful) < 80


def _markdown_heading_text(line: str) -> str | None:
    match = re.match(r"^(#{1,6})\s+(.+)$", line)
    if not match:
        return None
    return sanitize_text(match.group(2))


def _markdown_table_to_text(table_lines: list[str]) -> str:
    rows = []
    for line in table_lines:
        if re.fullmatch(r"\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*", line):
            continue
        cells = [sanitize_text(cell) for cell in line.strip("|").split("|")]
        cells = [cell for cell in cells if cell]
        if cells:
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _remove_html_chrome(soup: BeautifulSoup) -> None:
    for selector in [
        "script",
        "style",
        "nav",
        "footer",
        "iframe",
        "object",
        "embed",
        "form",
        ".footer",
        ".navigation",
        ".nav",
        ".print",
        ".email",
        ".printButton",
        ".emailButton",
        ".noprint",
        ".button",
        "#footer",
        "#navigation",
        "#nav",
    ]:
        for element in soup.select(selector):
            element.decompose()

    for element in soup.find_all(["a", "button", "input"]):
        text = sanitize_text(element.get_text(" ", strip=True))
        value = sanitize_text(str(element.get("value", "")))
        label = f"{text} {value}".strip().lower()
        href = str(element.get("href", "")).lower()
        if any(token in label or token in href for token in ["print", "email", "e-mail"]):
            element.decompose()


def _has_readable_text(root: Tag) -> bool:
    return bool(sanitize_text(root.get_text(" ", strip=True)))


def _extract_html_blocks(
    root: Tag,
    filename: str,
    title: str | None,
) -> tuple[list[TextBlock], str | None]:
    blocks: list[TextBlock] = []
    current_heading = _extract_primary_heading(root) or title

    if current_heading:
        blocks.append(
            TextBlock(
                text=current_heading,
                heading=current_heading,
                metadata={"source": filename, "style": "heading 1"},
            )
        )

    content_nodes = root.find_all(
        ["h1", "h2", "h3", "h4", "p", "li", "table", "div", "span"]
    )
    seen: set[str] = set()

    for node in content_nodes:
        if not isinstance(node, Tag) or _is_container_only_node(node):
            continue
        text = _node_to_text(node)
        if not text or text in seen:
            continue
        seen.add(text)

        tag_name = node.name.lower()
        classes = {str(item).lower() for item in node.get("class", [])}
        is_heading = tag_name in {"h1", "h2", "h3", "h4"} or "heading1" in classes
        if is_heading:
            current_heading = text
            blocks.append(
                TextBlock(
                    text=text,
                    heading=text,
                    metadata={"source": filename, "style": tag_name},
                )
            )
            continue

        blocks.append(
            TextBlock(
                text=text,
                heading=current_heading,
                metadata={"source": filename, "html_tag": tag_name},
            )
        )

    if len(blocks) <= 1:
        fallback_text = sanitize_text(root.get_text("\n\n", strip=True))
        if fallback_text and fallback_text != current_heading:
            blocks.append(
                TextBlock(
                    text=fallback_text,
                    heading=current_heading,
                    metadata={"source": filename, "html_tag": "body"},
                )
            )

    return blocks, current_heading


def _is_container_only_node(node: Tag) -> bool:
    if not node.name or node.name.lower() not in {"div", "span"}:
        return False
    nested_content = node.find(
        ["h1", "h2", "h3", "h4", "p", "li", "table", "div", "span"],
        recursive=False,
    )
    return nested_content is not None


def _extract_primary_heading(soup: BeautifulSoup | Tag) -> str | None:
    heading = soup.select_one("p.heading1") or soup.find("h1")
    if heading:
        return sanitize_text(heading.get_text(" ", strip=True))
    return None


def _extract_breadcrumbs(soup: BeautifulSoup) -> list[str]:
    breadcrumbs = soup.select_one("div.breadcrumbs")
    if not breadcrumbs:
        return []
    return [
        text
        for text in (_clean_text(item) for item in breadcrumbs.stripped_strings)
        if text and text not in {">", "/", "|"}
    ]


def _extract_links(root: Tag, filename: str) -> list[dict[str, str]]:
    links: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for anchor in root.find_all("a"):
        href = str(anchor.get("href", "")).strip()
        text = _clean_text(anchor.get_text(" ", strip=True))
        if not href or href.startswith(("javascript:", "mailto:")):
            continue
        if any(token in f"{text} {href}".lower() for token in ["print", "email", "e-mail"]):
            continue
        normalized = urljoin(filename, href)
        key = (text, normalized)
        if key in seen:
            continue
        seen.add(key)
        links.append({"text": text, "href": normalized})
    return links


def _extract_image_refs(root: Tag, filename: str) -> list[dict[str, str]]:
    images: list[dict[str, str]] = []
    seen: set[str] = set()
    for image in root.find_all("img"):
        src = str(image.get("src", "")).strip()
        if not src:
            continue
        normalized = urljoin(filename, src)
        if normalized in seen:
            continue
        seen.add(normalized)
        images.append(
            {
                "src": normalized,
                "filename": PurePosixPath(normalized).name,
                "alt": _clean_text(str(image.get("alt", ""))),
                "title": _clean_text(str(image.get("title", ""))),
            }
        )
    return images


def _extract_support_asset_refs(
    soup: BeautifulSoup,
    filename: str,
) -> dict[str, list[str]]:
    stylesheets = [
        urljoin(filename, str(link.get("href", "")).strip())
        for link in soup.find_all("link")
        if str(link.get("href", "")).strip().lower().endswith(".css")
    ]
    scripts = [
        urljoin(filename, str(script.get("src", "")).strip())
        for script in soup.find_all("script")
        if str(script.get("src", "")).strip().lower().endswith(".js")
    ]
    return {
        "stylesheets": sorted(set(stylesheets)),
        "scripts": sorted(set(scripts)),
    }


def _node_to_text(node: Tag) -> str:
    _replace_links_with_readable_text(node)
    if node.name and node.name.lower() == "li":
        return f"- {_clean_text(node.get_text(' ', strip=True))}"
    if node.name and node.name.lower() == "table":
        rows: list[str] = []
        for row in node.find_all("tr"):
            cells = [
                _clean_text(cell.get_text(" ", strip=True))
                for cell in row.find_all(["th", "td"])
            ]
            cells = [cell for cell in cells if cell]
            if cells:
                rows.append(" | ".join(cells))
        return "\n".join(rows)
    return _clean_text(node.get_text(" ", strip=True))


def _replace_links_with_readable_text(node: Tag) -> None:
    for anchor in node.find_all("a"):
        label = _clean_text(anchor.get_text(" ", strip=True))
        href = str(anchor.get("href", "")).strip()
        if label and href and not href.startswith(("javascript:", "mailto:")):
            anchor.string = f"{label} ({href})"


def _clean_text(value: str) -> str:
    return " ".join(sanitize_text(value).split())


def _extract_pdf_page_images(page, page_number: int) -> list[dict[str, str | int]]:
    images: list[dict[str, str | int]] = []
    try:
        for image_index, image in enumerate(getattr(page, "images", []) or []):
            name = str(getattr(image, "name", f"image-{image_index}"))
            images.append({"page": page_number, "name": name, "index": image_index})
    except Exception as exc:
        logger.warning(
            "Skipping PDF image metadata extraction on page %s: %s",
            page_number,
            exc,
        )
    return images


def _extract_printable_text(content: bytes) -> str:
    text = content.decode("utf-8", errors="ignore")
    if not text.strip():
        text = content.decode("latin-1", errors="ignore")
    cleaned = sanitize_text(text)
    lines = [" ".join(line.split()) for line in cleaned.splitlines()]
    lines = [line for line in lines if len(line) > 2]
    return "\n\n".join(lines)
